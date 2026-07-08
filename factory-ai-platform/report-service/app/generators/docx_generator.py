import io
from datetime import datetime
from typing import Any, Dict, List
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH


def generate_docx(
    title: str,
    summary: str,
    kpis: List[Dict[str, Any]],
    chart_data: List[Dict[str, Any]],
    top_alarms: List[Dict[str, Any]],
) -> bytes:
    doc = Document()

    # Title
    heading = doc.add_heading(title, 0)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = heading.runs[0]
    run.font.color.rgb = RGBColor(0x1C, 0x64, 0xF2)

    doc.add_paragraph(f"Ngày xuất: {datetime.now().strftime('%Y-%m-%d %H:%M')}")
    doc.add_paragraph()

    # KPI table
    if kpis:
        doc.add_heading("1. Chỉ số hiệu suất (KPI)", level=1)
        table = doc.add_table(rows=1, cols=2)
        table.style = "Light Grid Accent 1"
        hdr = table.rows[0].cells
        hdr[0].text = "Chỉ số"
        hdr[1].text = "Giá trị"
        for kpi in kpis:
            row = table.add_row().cells
            row[0].text = str(kpi.get("label", ""))
            row[1].text = str(kpi.get("value", ""))
        doc.add_paragraph()

    # Summary / report body
    doc.add_heading("2. Phân tích & Nhận xét", level=1)
    for line in summary.split("\n"):
        clean = line.strip().lstrip("#").lstrip("*").strip()
        if clean:
            doc.add_paragraph(clean)
    doc.add_paragraph()

    # Hourly / daily production table
    if chart_data:
        doc.add_heading("3. Sản lượng theo thời gian", level=1)
        keys = list(chart_data[0].keys()) if chart_data else []
        tbl = doc.add_table(rows=1, cols=len(keys))
        tbl.style = "Light List Accent 1"
        for i, k in enumerate(keys):
            tbl.rows[0].cells[i].text = k.upper()
        for row_data in chart_data:
            row = tbl.add_row().cells
            for i, k in enumerate(keys):
                row[i].text = str(row_data.get(k, ""))
        doc.add_paragraph()

    # Active alarms
    if top_alarms:
        doc.add_heading("4. Cảnh báo đang hoạt động", level=1)
        for alarm in top_alarms:
            doc.add_paragraph(
                f"[{alarm.get('severity', '')}] {alarm.get('machineName', '')} — {alarm.get('message', '')}",
                style="List Bullet",
            )

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
